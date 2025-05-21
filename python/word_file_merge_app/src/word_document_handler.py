# /src/word_document_handler.py

# system library -------------------------------------------------------------
import os
from io import BytesIO
from datetime import datetime

# packages -------------------------------------------------------------------
import docx
import docx.shared

# local library --------------------------------------------------------------
from config.document_style      import DocumentStyle
from src.sql_handler            import SqlHandler
from query.generated_file       import QueryGeneratedFile

class WordDocumentHandler(object):
    def __init__(self, template_id:int, generated_file_name:str ):
        self.document:docx.Document()       = None
        self.document_bytes_io              = None
        self.document_raw                   = None
        self.generated_file_id:int          = 0
        self.generated_file_name            = generated_file_name
        
        #self.output_file_prefix:str         = output_file_prefix
        
        self.output_file_word               = None
        self.output_path_word               = None
        self.query_handler:SqlHandler       = SqlHandler()
        self.template_id:int                = 0

        self.__set_path()
        
    def __append_bullet_item(self, bullet_para, bullet_list_item:str):
        new_paragraph = self.document.add_paragraph(bullet_list_item,
                                                    style=DocumentStyle.BULLET_STYLE.value)
        new_paragraph.paragraph_format.line_spacing     = 1
        new_paragraph.paragraph_format.space_after      = 0
        new_paragraph.style.font.name                   = DocumentStyle.BULLET_FONT_NAME.value
        new_paragraph.style.font.size                   = docx.shared.Pt(DocumentStyle.BULLET_FONT_POINT.value)

        bullet_para.addnext(new_paragraph._element)

    def __format_placeholder(self, placeholder:str) -> str:
        new_ph = ''.join(['{{',placeholder,'}}'])
        return(new_ph)

    def __get_document_bytes(self):
        self.document_bytes_io = BytesIO()
        self.document.save(self.document_bytes_io)

    def __load_document(self):
        """
        Load document from the database
        """
        doc_stream = BytesIO(self.document_raw)
        self.document = docx.Document(doc_stream)

    def __save_document_database(self):
        """
        Save the file to the database: available for future download
        """
        qry = QueryGeneratedFile.FILE_INSERT
        # convert for saving
        self.__get_document_bytes()

        # insert the generated file into the database
        self.query_handler.insert_query(qry,
                                        1, 
                                        self.generated_file_name,
                                        'docx',
                                        self.document_bytes_io.getvalue())
        
        self.generated_file_id = self.query_handler.GetLastId()
        qry = None


    def __set_path(self):
        self.output_path_word = os.path.join(DocumentStyle.OUTPUT_DIRECTORY_WORD.value,
                                             self.generated_file_name)


    def generate_word_document(self, template_file, placeholder_data, replacement_data):
        """
        Generate the word document replacing all placeholders with user selected data
        """
        # template_file
        self.document_raw = template_file
        self.__load_document()
        
        # replace all the placeholders based on type        
        for placeholder in placeholder_data:
           
            current_placeholder = placeholder['placeholder_name']
            current_replacement = replacement_data[current_placeholder]

            # depending on type of placeholder run the applicable action
            match placeholder['placeholder_type']:
                case "static":
                    self.replace_word(placeholder=current_placeholder, replace_with=current_replacement)
                case "inline":
                    self.replace_word(placeholder=current_placeholder, replace_with=current_replacement)
                case "list":
                    self.replace_paragraph_list(placeholder=current_placeholder, replace_with=current_replacement, separator=', ')
                case "bullet":
                    self.add_bullet_list(placeholder=current_placeholder, bullet_items=current_replacement)
                case _:
                    pass

        # save to database
        self.__save_document_database()

        return 0


    def add_bullet_list(self, placeholder:str, bullet_items:list):
        """
        Add the replacement text as bullet items
        """
        for current_paragraph in self.document.paragraphs:
            if placeholder in current_paragraph.text:
                paragraph_element   = current_paragraph._element
                paragraph_parent    = paragraph_element.getparent()
                # reverse because the append pops the list items
                bullet_items.reverse()
                for bullet_item in bullet_items:
                    self.__append_bullet_item(bullet_para=current_paragraph._element, bullet_list_item=bullet_item['item_text'])
                # now remove the placeholder bullet item
                paragraph_parent.remove(paragraph_element)


    def replace_word(self, placeholder:str, replace_with:dict):
        """Replacing a specific word with selected items in UI"""
        for paragraph in self.document.paragraphs:
            if placeholder in paragraph.text:
                replacement_text = ''.join(replace_with[0]['item_text'])
                paragraph.text = paragraph.text.replace(self.__format_placeholder(placeholder), replacement_text)
        
    
    def replace_paragraph_list(self, placeholder:str, replace_with:list, separator:str =''):
        """Replacing a specific word with selected items in UI"""
        # replacements = []
        # for list_item in replace_with:
        #     replacements.append(list_item['item_text'])
        
        replacements = [list_item['item_text'] for list_item in replace_with]

        for paragraph in self.document.paragraphs:
            if placeholder in paragraph.text:
                match len(replacements):
                    case 0:
                        replacement_text = ' '
                    case 1:
                        replacement_text = replacements[0]
                    case 2:
                        replacement_text = '{} and {}'.format(''.join(replacements[:-1]), replacements[-1])
                    case _:
                        replacement_text = '{} and {}'.format(separator.join(replacements[:-1]), replacements[-1])

                paragraph.text = paragraph.text.replace(self.__format_placeholder(placeholder), replacement_text)
    
